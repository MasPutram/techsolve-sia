# -*- coding: utf-8 -*-
"""
generate_artikel.py
Generator artikel ilmiah (.docx) untuk UAS Sistem Informasi Akuntansi.

Ketentuan dosen yang dipatuhi:
- Maksimal 5 halaman
- Times New Roman 10pt, spasi 1 (single), margin normal 2.54cm
- Struktur: Judul, Abstrak, Pendahuluan, Tinjauan Pustaka, Metode, Hasil, Penutup
- Judul section: bold 12pt; tabel TNR 9pt berbingkai; code SQL monospace 8.5pt

Cara pakai:
    pip install -r requirements.txt
    python generate_artikel.py
Output: artikel_ilmiah_sia.docx (di folder yang sama)

Catatan sitasi:
    Penanda [SITASI-N] disebar di body. Daftar topiknya ada di bagian
    "Daftar Pustaka". Ganti penanda dengan sitasi Mendeley saat finalisasi.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ------------------------------------------------------------------
# Konstanta gaya
# ------------------------------------------------------------------
FONT_BODY = "Times New Roman"
FONT_MONO = "Consolas"
SIZE_BODY = 10
SIZE_TITLE = 14
SIZE_HEADING = 12
SIZE_TABLE = 9
SIZE_CODE = 8.5

OUTPUT = "artikel_ilmiah_sia.docx"


# ------------------------------------------------------------------
# Helper format
# ------------------------------------------------------------------
def set_single_spacing(paragraph, space_after=4):
    """Atur spasi 1 (single) dan jarak antar paragraf kecil."""
    pf = paragraph.paragraph_format
    pf.line_spacing = 1.0
    pf.space_before = Pt(0)
    pf.space_after = Pt(space_after)


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_single_spacing(p, space_after=6)
    r = p.add_run(text)
    r.bold = True
    r.font.name = FONT_BODY
    r.font.size = Pt(SIZE_TITLE)
    return p


def add_center_line(doc, text, italic=False, size=SIZE_BODY, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_single_spacing(p, space_after=2)
    r = p.add_run(text)
    r.italic = italic
    r.bold = bold
    r.font.name = FONT_BODY
    r.font.size = Pt(size)
    return p


def add_heading(doc, text):
    p = doc.add_paragraph()
    set_single_spacing(p, space_after=4)
    p.paragraph_format.space_before = Pt(8)
    r = p.add_run(text)
    r.bold = True
    r.font.name = FONT_BODY
    r.font.size = Pt(SIZE_HEADING)
    return p


def add_body(doc, text, justify=True, bold=False, italic=False):
    p = doc.add_paragraph()
    set_single_spacing(p)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0.75)
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.name = FONT_BODY
    r.font.size = Pt(SIZE_BODY)
    return p


def add_plain(doc, text, justify=True):
    """Paragraf tanpa indent (untuk daftar/abstrak)."""
    p = doc.add_paragraph()
    set_single_spacing(p)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.font.name = FONT_BODY
    r.font.size = Pt(SIZE_BODY)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_single_spacing(p, space_after=6)
    r = p.add_run(text)
    r.italic = True
    r.font.name = FONT_BODY
    r.font.size = Pt(9)
    return p


def add_image_placeholder(doc, caption):
    """Kotak placeholder gambar + caption (gambar diinsert manual)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_single_spacing(p, space_after=2)
    r = p.add_run("[ TEMPATKAN GAMBAR DI SINI ]")
    r.font.name = FONT_BODY
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    add_caption(doc, caption)


def add_image_grid(doc, captions):
    """Susun placeholder gambar dalam grid 2 kolom agar compact.
    captions: list caption; tiap sel berisi kotak placeholder + caption."""
    n = len(captions)
    rows = (n + 1) // 2
    table = doc.add_table(rows=rows, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, cap in enumerate(captions):
        cell = table.rows[idx // 2].cells[idx % 2]
        # paragraf placeholder
        p0 = cell.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_single_spacing(p0, space_after=1)
        r0 = p0.add_run("[ TEMPATKAN GAMBAR ]")
        r0.font.name = FONT_BODY
        r0.font.size = Pt(8)
        r0.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        # paragraf caption
        pc = cell.add_paragraph()
        pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_single_spacing(pc, space_after=1)
        rc = pc.add_run(cap)
        rc.italic = True
        rc.font.name = FONT_BODY
        rc.font.size = Pt(8.5)
    # jika jumlah ganjil, kosongkan sel terakhir
    sp = doc.add_paragraph()
    set_single_spacing(sp, space_after=2)


def add_code_block(doc, code):
    """Blok kode monospace kecil dengan spasi rapat."""
    for line in code.strip("\n").split("\n"):
        p = doc.add_paragraph()
        set_single_spacing(p, space_after=0)
        p.paragraph_format.left_indent = Cm(0.5)
        r = p.add_run(line if line else " ")
        r.font.name = FONT_MONO
        # set juga untuk east-asian agar konsisten
        r._element.rPr.rFonts.set(qn("w:cs"), FONT_MONO)
        r.font.size = Pt(SIZE_CODE)


def style_table(table):
    """Terapkan font TNR kecil pada seluruh sel tabel."""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                set_single_spacing(p, space_after=0)
                for r in p.runs:
                    r.font.name = FONT_BODY
                    r.font.size = Pt(SIZE_TABLE)


def add_table(doc, header, rows, header_bold=True):
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = header_bold
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            cells[i].paragraphs[0].add_run(str(val))
    style_table(table)
    # paragraf kecil setelah tabel
    sp = doc.add_paragraph()
    set_single_spacing(sp, space_after=2)
    return table


# ------------------------------------------------------------------
# Bangun dokumen
# ------------------------------------------------------------------
def build():
    doc = Document()

    # Margin normal 2.54cm + font default
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Cm(2.54)
    sec.left_margin = sec.right_margin = Cm(2.54)
    normal = doc.styles["Normal"]
    normal.font.name = FONT_BODY
    normal.font.size = Pt(SIZE_BODY)
    normal.paragraph_format.line_spacing = 1.0
    normal.paragraph_format.space_after = Pt(4)

    # ---------- JUDUL ----------
    add_title(
        doc,
        "Implementasi Sistem Informasi Akuntansi Berbasis Web dengan "
        "Otomatisasi Penjurnalan Menggunakan Database Trigger "
        "pada UMKM Jasa Konsultan IT (Studi Kasus: TechSolve Consulting)",
    )
    add_center_line(doc, "[Nama Mahasiswa] — [NIM]", bold=True)
    add_center_line(
        doc,
        "Program Studi Informatika, Fakultas Teknik dan Sains, "
        "Universitas Muhammadiyah Purwokerto",
        size=9,
    )
    add_center_line(doc, "e-mail: [email kamu]", size=9)
    doc.add_paragraph()

    # ---------- ABSTRAK ----------
    add_heading(doc, "Abstrak")
    add_plain(
        doc,
        "Pencatatan keuangan secara manual pada usaha mikro, kecil, dan menengah "
        "(UMKM) jasa rentan terhadap kesalahan manusia (human error), khususnya "
        "pada proses penjurnalan dan penyusunan laporan keuangan yang sering "
        "tidak sinkron dengan data transaksi. Penelitian ini bertujuan membangun "
        "Sistem Informasi Akuntansi (SIA) berbasis web yang mengotomatiskan "
        "proses penjurnalan menggunakan database trigger. Sistem dikembangkan "
        "dengan framework Next.js 14, basis data MySQL 8, dan driver mysql2 tanpa "
        "Object-Relational Mapping agar logika trigger transparan di level basis "
        "data. Tiga trigger (AFTER INSERT, AFTER UPDATE, dan AFTER DELETE) "
        "dipasang pada tabel detail transaksi sehingga setiap perubahan data "
        "transaksi otomatis menjaga konsistensi buku jurnal. Validasi double-entry "
        "(total debit sama dengan total kredit) diterapkan pada sisi antarmuka "
        "maupun server. Hasil pengujian menunjukkan bahwa 7 transaksi contoh "
        "menghasilkan 14 entri jurnal secara otomatis dengan total debit dan "
        "kredit seimbang sebesar Rp85.500.000, laba bersih Rp3.500.000, serta "
        "neraca yang seimbang sebesar Rp52.500.000. Dengan demikian, pemanfaatan "
        "database trigger terbukti menghilangkan penjurnalan manual dan menjamin "
        "akurasi laporan keuangan.",
    )
    kk = doc.add_paragraph()
    set_single_spacing(kk)
    rk = kk.add_run("Kata kunci: ")
    rk.bold = True
    rk.font.name = FONT_BODY
    rk.font.size = Pt(SIZE_BODY)
    rk2 = kk.add_run(
        "sistem informasi akuntansi; database trigger; double-entry; "
        "otomatisasi penjurnalan; Next.js"
    )
    rk2.italic = True
    rk2.font.name = FONT_BODY
    rk2.font.size = Pt(SIZE_BODY)

    # ---------- 1. PENDAHULUAN ----------
    add_heading(doc, "1. PENDAHULUAN")
    add_body(
        doc,
        "Sistem Informasi Akuntansi (SIA) merupakan komponen penting dalam "
        "pengelolaan keuangan suatu entitas usaha karena berfungsi mengumpulkan, "
        "mencatat, dan mengolah data transaksi menjadi informasi keuangan yang "
        "berguna bagi pengambilan keputusan [SITASI-1]. Bagi pelaku UMKM, "
        "khususnya pada sektor jasa, ketersediaan informasi keuangan yang akurat "
        "dan tepat waktu menjadi kebutuhan yang semakin mendesak seiring "
        "meningkatnya volume transaksi [SITASI-2].",
    )
    add_body(
        doc,
        "Namun, banyak UMKM masih melakukan pencatatan secara manual menggunakan "
        "buku atau lembar kerja sederhana. Praktik ini rawan terhadap kesalahan "
        "manusia (human error), seperti salah menempatkan akun, ketidakseimbangan "
        "antara debit dan kredit, serta laporan keuangan yang tidak sinkron "
        "dengan jurnal [SITASI-3]. Kesalahan tersebut berdampak langsung pada "
        "kualitas laporan laba rugi, perubahan modal, dan neraca.",
    )
    add_body(
        doc,
        "Untuk mengatasi permasalahan tersebut, penelitian ini membangun SIA "
        "berbasis web yang mengotomatiskan penjurnalan pada level basis data "
        "menggunakan mekanisme database trigger, dengan studi kasus TechSolve "
        "Consulting (UMKM fiktif jasa konsultan teknologi informasi). Setiap "
        "transaksi yang dimasukkan otomatis membentuk entri buku jurnal tanpa "
        "intervensi manual. Tujuan penelitian adalah merancang dan menguji "
        "sistem tersebut dalam menjaga prinsip double-entry serta menghasilkan "
        "laporan keuangan yang seimbang, sekaligus menyediakan rujukan praktis "
        "penerapan trigger untuk menjamin integritas data akuntansi pada UMKM.",
    )

    # ---------- 2. TINJAUAN PUSTAKA ----------
    add_heading(doc, "2. TINJAUAN PUSTAKA")
    add_body(
        doc,
        "Sistem Informasi Akuntansi didefinisikan sebagai kumpulan sumber daya "
        "yang dirancang untuk mengubah data keuangan menjadi informasi yang "
        "dibutuhkan oleh berbagai pihak [SITASI-1]. SIA modern umumnya "
        "terkomputerisasi sehingga mampu memproses transaksi dalam jumlah besar "
        "dengan tingkat akurasi yang tinggi [SITASI-2].",
    )
    add_body(
        doc,
        "Pencatatan berpasangan (double-entry bookkeeping) merupakan prinsip "
        "dasar akuntansi yang menyatakan bahwa setiap transaksi dicatat pada "
        "minimal dua akun dengan nilai total debit yang sama dengan total kredit "
        "[SITASI-4]. Prinsip ini menjadi mekanisme kontrol bawaan untuk menjaga "
        "keseimbangan persamaan akuntansi, yaitu Aset = Kewajiban + Ekuitas.",
    )
    add_body(
        doc,
        "Database trigger adalah prosedur khusus yang tersimpan dan dieksekusi "
        "secara otomatis oleh sistem manajemen basis data ketika terjadi "
        "peristiwa tertentu pada sebuah tabel [SITASI-5]. Pada MySQL, trigger "
        "dapat diatur untuk berjalan sebelum (BEFORE) atau sesudah (AFTER) "
        "operasi INSERT, UPDATE, maupun DELETE. Trigger banyak dimanfaatkan untuk "
        "menjaga integritas data dan mengotomatiskan proses turunan, seperti "
        "pembentukan jurnal dari data transaksi.",
    )
    add_body(
        doc,
        "Next.js merupakan kerangka kerja berbasis React yang mendukung "
        "perenderan sisi server dan penyediaan API dalam satu basis kode "
        "[SITASI-6], sehingga memisahkan lapisan antarmuka, logika bisnis, dan "
        "penyimpanan data secara ringkas. Penelitian terdahulu telah membahas "
        "pengembangan aplikasi akuntansi untuk UMKM [SITASI-7], otomatisasi "
        "pengolahan data transaksi [SITASI-8], serta pemanfaatan trigger untuk "
        "menjaga konsistensi data [SITASI-9]. Penelitian ini melengkapi kajian "
        "tersebut dengan menekankan otomatisasi penjurnalan double-entry "
        "sepenuhnya pada level trigger MySQL.",
    )

    # ---------- 3. METODE ----------
    add_heading(doc, "3. METODE PENELITIAN")
    add_body(
        doc,
        "Penelitian ini menggunakan pendekatan pengembangan perangkat lunak model "
        "Waterfall yang terdiri atas tahap analisis kebutuhan, perancangan, "
        "implementasi, dan pengujian. Objek studi kasus adalah TechSolve "
        "Consulting, UMKM jasa konsultan TI yang aktivitasnya meliputi instalasi "
        "jaringan, pemeliharaan, perbaikan perangkat, dan pelatihan karyawan.",
    )
    add_body(doc, "Teknologi yang digunakan dirangkum pada Tabel 1.", )
    add_caption(doc, "Tabel 1. Teknologi yang Digunakan")
    add_table(
        doc,
        ["Komponen", "Teknologi"],
        [
            ["Framework", "Next.js 14 (App Router)"],
            ["Bahasa", "TypeScript"],
            ["Basis Data", "MySQL 8.0"],
            ["Driver Basis Data", "mysql2/promise (raw query, tanpa ORM)"],
            ["Styling", "Tailwind CSS"],
            ["Validasi", "Zod"],
        ],
    )
    add_body(
        doc,
        "Arsitektur sistem mengikuti alur berlapis: pengguna berinteraksi melalui "
        "antarmuka (UI), permintaan diteruskan ke API Route, kemudian diolah ke "
        "basis data MySQL. Pada saat data detail transaksi disimpan, trigger "
        "MySQL otomatis menulis entri ke tabel jurnal, yang selanjutnya menjadi "
        "sumber tunggal bagi seluruh laporan keuangan. Alur tersebut dapat "
        "dituliskan sebagai: Pengguna → UI → API → MySQL → "
        "Trigger → Buku Jurnal → Laporan.",
    )
    add_body(
        doc,
        "Basis data dirancang dengan empat tabel utama (Tabel 2 dan Gambar 1): "
        "chart_of_accounts (master akun), transactions (header transaksi), "
        "transaction_details (baris debit/kredit), dan journal_entries (buku "
        "jurnal yang diisi otomatis oleh trigger). Kolom penghubung ref_detail_id "
        "ditambahkan pada tabel jurnal untuk menautkan setiap baris detail dengan "
        "tepat satu entri jurnal (relasi satu-ke-satu), sehingga operasi "
        "pembaruan dan penghapusan menjadi presisi.",
    )
    add_caption(doc, "Tabel 2. Struktur Tabel Basis Data")
    add_table(
        doc,
        ["Tabel", "Fungsi", "Kunci Relasi"],
        [
            ["chart_of_accounts", "Master akun (CoA)", "PK kode_akun"],
            ["transactions", "Header transaksi", "PK id"],
            ["transaction_details", "Baris debit/kredit", "FK transaction_id, kode_akun"],
            ["journal_entries", "Buku jurnal (diisi trigger)", "FK ref_transaction_id, ref_detail_id"],
        ],
    )
    add_image_placeholder(doc, "Gambar 1. Entity Relationship Diagram (ERD) Sistem")

    add_body(
        doc,
        "Inti otomatisasi terletak pada tiga trigger berikut yang dipasang pada "
        "tabel transaction_details. Kode trigger ditampilkan secara ringkas pada "
        "Kode Program 1.",
    )
    add_caption(doc, "Kode Program 1. Implementasi Trigger MySQL")
    add_code_block(
        doc,
        """
-- AFTER INSERT: detail baru -> buat entri jurnal otomatis
CREATE TRIGGER trg_jurnal_insert
AFTER INSERT ON transaction_details FOR EACH ROW
BEGIN
  DECLARE v_tgl DATE; DECLARE v_desk VARCHAR(255);
  SELECT tanggal, deskripsi INTO v_tgl, v_desk
    FROM transactions WHERE id = NEW.transaction_id;
  INSERT INTO journal_entries
    (tanggal, deskripsi, kode_akun, debit, kredit,
     ref_transaction_id, ref_detail_id)
  VALUES (v_tgl, v_desk, NEW.kode_akun, NEW.debit,
     NEW.kredit, NEW.transaction_id, NEW.id);
END;

-- AFTER UPDATE: detail berubah -> jurnal terkait diperbarui
CREATE TRIGGER trg_jurnal_update
AFTER UPDATE ON transaction_details FOR EACH ROW
BEGIN
  DECLARE v_tgl DATE; DECLARE v_desk VARCHAR(255);
  SELECT tanggal, deskripsi INTO v_tgl, v_desk
    FROM transactions WHERE id = NEW.transaction_id;
  UPDATE journal_entries
     SET tanggal=v_tgl, deskripsi=v_desk, kode_akun=NEW.kode_akun,
         debit=NEW.debit, kredit=NEW.kredit,
         ref_transaction_id=NEW.transaction_id
   WHERE ref_detail_id = OLD.id;
END;

-- AFTER DELETE: detail dihapus -> jurnal terkait dihapus
CREATE TRIGGER trg_jurnal_delete
AFTER DELETE ON transaction_details FOR EACH ROW
BEGIN
  DELETE FROM journal_entries WHERE ref_detail_id = OLD.id;
END;
""",
    )
    add_body(
        doc,
        "Validasi double-entry dilakukan dua lapis. Pada antarmuka, tombol simpan "
        "dinonaktifkan selama total debit belum sama dengan total kredit. Pada "
        "sisi server, pustaka Zod memverifikasi ulang keseimbangan sebelum data "
        "disimpan, sehingga transaksi yang tidak seimbang ditolak.",
    )

    # ---------- 4. HASIL DAN PEMBAHASAN ----------
    add_heading(doc, "4. HASIL DAN PEMBAHASAN")
    add_body(
        doc,
        "Sistem berhasil diimplementasikan dengan lima fitur utama, yaitu "
        "pengelolaan daftar akun (CRUD Chart of Accounts), input transaksi dengan "
        "validasi double-entry, jurnal umum, buku besar dengan saldo berjalan, "
        "serta tiga laporan keuangan otomatis. Cuplikan antarmuka fitur utama "
        "ditunjukkan pada Gambar 2 dan Gambar 3.",
    )
    add_image_grid(
        doc,
        [
            "Gambar 2. Form Input Transaksi (Validasi Double-Entry)",
            "Gambar 3. Laporan Neraca (Seimbang)",
        ],
    )

    add_body(
        doc,
        "Data uji terdiri atas 16 akun standar dan 7 transaksi contoh. Cuplikan "
        "transaksi disajikan pada Tabel 3. Setelah seluruh transaksi disimpan, "
        "tabel journal_entries terisi 14 baris secara otomatis tanpa satu pun "
        "perintah penyisipan manual ke tabel jurnal—membuktikan trigger "
        "AFTER INSERT bekerja.",
    )
    add_caption(doc, "Tabel 3. Cuplikan Transaksi Contoh")
    add_table(
        doc,
        ["Tgl", "Deskripsi", "Debit", "Kredit", "Jumlah (Rp)"],
        [
            ["01/06", "Setoran modal awal", "Kas", "Modal Pemilik", "50.000.000"],
            ["02/06", "Beli peralatan komputer", "Peralatan", "Kas", "20.000.000"],
            ["10/06", "Pendapatan jasa instalasi", "Kas", "Pendapatan Jasa", "8.000.000"],
            ["15/06", "Bayar gaji karyawan", "Beban Gaji", "Kas", "3.000.000"],
            ["25/06", "Pengambilan pribadi (Prive)", "Prive", "Kas", "1.000.000"],
        ],
    )
    add_body(
        doc,
        "Pengujian keandalan trigger terhadap tiga operasi data dirangkum pada "
        "Tabel 4. Penambahan transaksi seimbang menambah dua entri jurnal "
        "(14→16), pengubahan nominal memperbarui entri jurnal terkait, dan "
        "penghapusan mengembalikan jumlah entri ke kondisi semula (16→14). "
        "Transaksi dengan debit dan kredit tidak seimbang berhasil ditolak.",
    )
    add_caption(doc, "Tabel 4. Hasil Pengujian Trigger dan Validasi")
    add_table(
        doc,
        ["Operasi", "Aksi Uji", "Hasil"],
        [
            ["AFTER INSERT", "Tambah 1 transaksi seimbang", "Jurnal 14 → 16 (otomatis)"],
            ["AFTER UPDATE", "Ubah nominal Rp2 jt → Rp7 jt", "Nilai jurnal ikut berubah"],
            ["AFTER DELETE", "Hapus transaksi uji", "Jurnal 16 → 14 (sinkron)"],
            ["Validasi", "Debit ≠ Kredit", "Ditolak (HTTP 400)"],
        ],
    )
    add_body(
        doc,
        "Laporan keuangan yang dihasilkan dari data jurnal konsisten dan seimbang "
        "(Tabel 5). Total pendapatan Rp8.000.000 dikurangi total beban "
        "Rp4.500.000 menghasilkan laba bersih Rp3.500.000. Modal akhir "
        "Rp52.500.000 berasal dari modal awal Rp50.000.000 ditambah laba bersih "
        "dikurangi prive Rp1.000.000. Neraca menunjukkan total aset sama dengan "
        "kewajiban ditambah ekuitas sebesar Rp52.500.000, sehingga persamaan "
        "akuntansi terpenuhi.",
    )
    add_caption(doc, "Tabel 5. Ringkasan Hasil Laporan Keuangan")
    add_table(
        doc,
        ["Laporan", "Komponen", "Nilai (Rp)"],
        [
            ["Laba Rugi", "Total Pendapatan", "8.000.000"],
            ["Laba Rugi", "Total Beban", "4.500.000"],
            ["Laba Rugi", "Laba Bersih", "3.500.000"],
            ["Perubahan Modal", "Modal Akhir", "52.500.000"],
            ["Neraca", "Total Aset = Pasiva", "52.500.000 (seimbang)"],
        ],
    )
    add_image_placeholder(
        doc, "Gambar 4. Output SHOW TRIGGERS pada MySQL Workbench (Bukti Trigger)"
    )

    # ---------- 5. PENUTUP ----------
    add_heading(doc, "5. PENUTUP")
    add_body(
        doc,
        "Berdasarkan hasil implementasi dan pengujian, dapat disimpulkan beberapa "
        "hal. Pertama, otomatisasi penjurnalan menggunakan database trigger MySQL "
        "berhasil menghilangkan kebutuhan input jurnal manual dan menjamin "
        "sinkronisasi antara data transaksi dan buku jurnal pada operasi INSERT, "
        "UPDATE, maupun DELETE. Kedua, penerapan validasi double-entry pada sisi "
        "antarmuka dan server menjaga keseimbangan debit dan kredit pada setiap "
        "transaksi. Ketiga, laporan keuangan yang dihitung langsung dari tabel "
        "jurnal terbukti akurat dan seimbang, ditunjukkan oleh neraca yang balance "
        "sebesar Rp52.500.000.",
    )
    add_body(
        doc,
        "Penelitian ini dapat dikembangkan lebih lanjut dengan menambahkan "
        "autentikasi multi-pengguna, fitur tutup buku per periode akuntansi, "
        "ekspor laporan ke format PDF, serta jejak audit (audit trail) untuk "
        "meningkatkan akuntabilitas sistem.",
    )

    # ---------- DAFTAR PUSTAKA ----------
    add_heading(doc, "DAFTAR PUSTAKA")

    # Pemetaan penanda [SITASI-N] -> referensi (untuk diisi via Mendeley)
    pmap = doc.add_paragraph()
    set_single_spacing(pmap, space_after=4)
    pm = pmap.add_run(
        "Pemetaan penanda sitasi (isi sitasi dalam teks via Mendeley sesuai peta "
        "ini): [SITASI-1] dan [SITASI-7] -> Firdaus & Widyasastrena (2017); "
        "[SITASI-2] -> Nugraha dkk. (2023); [SITASI-3] dan [SITASI-8] -> "
        "Ramadani dkk. (2025); [SITASI-4] -> Prasetyo (2013); "
        "[SITASI-5] dan [SITASI-9] -> Aidjili (2025); [SITASI-6] -> "
        "Maharani (2025) dan Naresvari & Susetyo (2025)."
    )
    pm.italic = True
    pm.font.name = FONT_BODY
    pm.font.size = Pt(9)

    # Daftar referensi (terverifikasi; urut alfabetis)
    referensi = [
        "Aidjili, M. (2025). Implementasi trigger dan view untuk mendukung "
        "konsistensi dan efisiensi pengolahan data pada sistem database (Studi "
        "kasus: Toko Nanda Pekalongan). Jurnal Komputer, Informasi dan "
        "Teknologi, 5(2). https://doi.org/10.53697/jkomitek.v5i2.3940",
        "Firdaus, D. W., & Widyasastrena, D. (2017). Perancangan sistem "
        "informasi akuntansi koperasi dan UMKM berbasis technopreneur. Jurnal "
        "Riset Akuntansi dan Keuangan, 5(2), 1423-1440. "
        "https://doi.org/10.17509/jrak.v5i2.8124",
        "Maharani, P. (2025). Pengembangan website PT. Rantangin Digital "
        "Indonesia menggunakan framework Next.js dan Tailwind CSS. Repeater: "
        "Publikasi Teknik Informatika dan Jaringan, 3(1), 129-137. "
        "https://doi.org/10.62951/repeater.v3i1.355",
        "Naresvari, E., & Susetyo, Y. A. (2025). Penerapan JavaScript React pada "
        "perancangan front-end website UMKM Jemari Ragil. IT-Explore: Jurnal "
        "Penerapan Teknologi Informasi dan Komunikasi, 4(1), 16-32. "
        "https://doi.org/10.24246/itexplore.v4i1.2025.pp16-32",
        "Nugraha, N., Budiyono, I., Nurhayati, I., & Arumsari, V. (2023). "
        "Pemanfaatan sistem informasi akuntansi pada UMKM di Kota Semarang. "
        "KEUNIS, 11(1), 95-104. https://doi.org/10.32497/keunis.v11i1.4079",
        "Prasetyo, W. (2013). Membongkar akuntansi double entry systems. Jurnal "
        "Akuntansi Multiparadigma, 4(2), 308-321. "
        "https://doi.org/10.18202/jamal.2013.08.7199",
        "Ramadani, F., Budianto, F., Saputra, M. D., & Faizah, Y. L. N. (2025). "
        "Penerapan sistem informasi akuntansi berbasis web pada UMKM (Studi "
        "kasus: Diska Beauty Salon). DEVICE: Journal of Information System, "
        "Computer Science and Information Technology, 6(2). "
        "https://doi.org/10.46576/device.v6i2.7206",
    ]
    for ref in referensi:
        p = doc.add_paragraph()
        set_single_spacing(p, space_after=2)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        # indentasi gantung (hanging indent) gaya daftar pustaka
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.first_line_indent = Cm(-0.75)
        r = p.add_run(ref)
        r.font.name = FONT_BODY
        r.font.size = Pt(9)

    doc.save(OUTPUT)
    print(f"Berhasil membuat: {OUTPUT}")


if __name__ == "__main__":
    build()
